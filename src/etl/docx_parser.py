from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from src.etl.base import detect_case_from_path, normalize_label
from src.models.schemas import ReferenceHypothesis, SourceRef, TextChunk


def _strip_number_prefix(text: str) -> tuple[int | None, str]:
    match = re.match(r"^(\d+)\.\s*(.+)$", text.strip(), re.DOTALL)
    if match:
        return int(match.group(1)), normalize_label(match.group(2))
    return None, normalize_label(text)


def parse_docx_hypotheses(path: Path, case_id: str | None = None) -> list[ReferenceHypothesis]:
    """Extract organizer example hypotheses (format reference only, not graph/RAG)."""
    path = Path(path)
    case = detect_case_from_path(path)
    if case_id is None:
        case_id = case[0] if case else path.stem.lower().replace(" ", "_")

    hypotheses: list[ReferenceHypothesis] = []

    doc = Document(path)
    source = SourceRef(file=path.name)

    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            cells = [normalize_label(cell.text) for cell in row.cells if normalize_label(cell.text)]
            if not cells:
                continue

            raw_title = cells[0]
            index, title = _strip_number_prefix(raw_title)
            if not title:
                continue
            if index is None:
                index = len(hypotheses) + 1

            row_source = source.model_copy(
                update={"row": row_idx + 1, "fragment": title, "sheet": f"table_{table_idx + 1}"}
            )
            hypotheses.append(
                ReferenceHypothesis(index=index, title=title, case_id=case_id, source=row_source)
            )

    return hypotheses


def parse_docx_text(
    path: Path,
    *,
    case_id: str | None = None,
    chunk_type: str = "text",
    chunk_size: int = 1200,
    chunk_overlap: int = 150,
) -> list[TextChunk]:
    """Extract paragraph text from docx, split into RAG-sized chunks."""
    path = Path(path)
    case = detect_case_from_path(path)
    if case_id is None:
        case_id = case[0] if case else None

    paragraphs = [normalize_label(p.text) for p in Document(path).paragraphs if normalize_label(p.text)]
    if not paragraphs:
        return []

    full_text = "\n\n".join(paragraphs)
    chunks: list[TextChunk] = []
    start = 0
    chunk_num = 0

    while start < len(full_text):
        end = min(start + chunk_size, len(full_text))
        if end < len(full_text):
            split_at = full_text.rfind("\n", start, end)
            if split_at <= start:
                split_at = full_text.rfind(" ", start, end)
            if split_at > start:
                end = split_at

        text = full_text[start:end].strip()
        if text:
            chunk_num += 1
            chunks.append(
                TextChunk(
                    chunk_id=f"{path.stem}_{chunk_num}",
                    text=text,
                    source=SourceRef(file=path.name, fragment=text[:120]),
                    case_id=case_id,
                    chunk_type=chunk_type,
                    metadata={"chunk_index": chunk_num, "char_start": start},
                )
            )

        if end >= len(full_text):
            break
        start = max(end - chunk_overlap, start + 1)

    return chunks
