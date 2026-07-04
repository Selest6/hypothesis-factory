from __future__ import annotations

import csv
import io
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from src.feedback.store import save_feedback
from src.models.schemas import GeneratedHypothesis, PipelineResult, SourceRef

FONT_PATH = Path(__file__).resolve().parents[2] / "assets" / "fonts" / "Arial.ttf"
PDF_FONT_NAME = "ReportArial"


def _normalize_flow_text(text: str) -> str:
    """Склеить переносы из OCR/PDF в обычный текст для переноса по словам."""
    if not text:
        return ""
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    # дефисный перенос: доизме-\nльчение -> доизмельчение
    normalized = re.sub(r"(\S)-\n(\S)", r"\1\2", normalized)
    # остальные переносы строк (обложка книги и т.п.) -> пробел
    normalized = re.sub(r"\s*\n+\s*", " ", normalized)
    return re.sub(r" +", " ", normalized).strip()


def _register_pdf_font() -> str:
    if PDF_FONT_NAME in pdfmetrics.getRegisteredFontNames():
        return PDF_FONT_NAME
    if FONT_PATH.exists():
        pdfmetrics.registerFont(TTFont(PDF_FONT_NAME, str(FONT_PATH)))
        return PDF_FONT_NAME
    return "Helvetica"


def _pdf_styles(font_name: str) -> dict[str, ParagraphStyle]:
    body = ParagraphStyle(
        "PdfBody",
        fontName=font_name,
        fontSize=10,
        leading=14,
        alignment=TA_LEFT,
        spaceAfter=4,
    )
    return {
        "body": body,
        "title": ParagraphStyle(
            "PdfTitle",
            parent=body,
            fontSize=16,
            leading=20,
            spaceAfter=8,
        ),
        "heading": ParagraphStyle(
            "PdfHeading",
            parent=body,
            fontSize=12,
            leading=16,
            spaceBefore=10,
            spaceAfter=6,
        ),
        "meta": ParagraphStyle(
            "PdfMeta",
            parent=body,
            fontSize=10,
            leading=13,
            spaceAfter=3,
        ),
        "statement": ParagraphStyle(
            "PdfStatement",
            parent=body,
            fontSize=10,
            leading=14,
            leftIndent=6,
            spaceAfter=6,
        ),
    }


def _pdf_paragraph(text: str, style: ParagraphStyle) -> Paragraph:
    safe = xml_escape(_normalize_flow_text(text))
    return Paragraph(safe, style)


def _format_source(src: SourceRef | dict[str, Any]) -> str:
    if isinstance(src, SourceRef):
        data = src.model_dump()
    else:
        data = src
    parts = [str(data.get("file") or "—")]
    if data.get("sheet"):
        parts.append(f"лист {data['sheet']}")
    if data.get("row"):
        parts.append(f"строка {data['row']}")
    if data.get("page"):
        parts.append(f"стр. {data['page']}")
    line = ", ".join(parts)
    if data.get("fragment"):
        frag = _normalize_flow_text(str(data["fragment"]))
        if len(frag) > 200:
            frag = frag[:200] + "…"
        line += f" — {frag}"
    return line


def _format_risks(h: GeneratedHypothesis) -> tuple[str, str]:
    risks = h.risks
    if isinstance(risks, dict):
        return str(risks.get("technical", "—")), str(risks.get("economic", "—"))
    if isinstance(risks, list):
        tech = risks[0] if risks else "—"
        econ = risks[1] if len(risks) > 1 else "—"
        return str(tech), str(econ)
    return "—", "—"


def result_to_markdown(result: PipelineResult, constraints: str = "") -> str:
    lines = [
        f"# Фабрика гипотез — {result.case_name}",
        "",
        f"**KPI:** {result.kpi_goal}",
        f"**Ограничения:** {constraints or '—'}",
        f"**Режим:** {result.mode}",
        f"**Дата:** {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        "",
    ]
    for i, h in enumerate(result.hypotheses, 1):
        lines.extend(_hypothesis_md_block(h, i))
    return "\n".join(lines)


def _hypothesis_md_block(h: GeneratedHypothesis, index: int) -> list[str]:
    lines = [f"## {index}. {h.title}", "", f"**Формулировка:** {h.full_statement}", ""]
    if h.mechanism:
        lines.extend([f"**Механизм:** {h.mechanism}", ""])
    if h.kpi_impact:
        lines.extend([f"**Влияние на KPI:** {h.kpi_impact}", ""])
    if h.scores:
        s = h.scores
        lines.append(
            f"**Оценки:** итого {s.total:.2f} | новизна {s.novelty:.2f} | "
            f"обоснованность {s.groundedness:.2f} | ценность {s.value:.2f} | риск {s.risk:.2f}"
        )
        lines.append("")
    if h.prior_art_snippet:
        sim = (h.prior_art_similarity or 0) * 100
        lines.append(f"**Ближайший фрагмент литературы:** «{h.prior_art_snippet}» (сходство {sim:.0f}%)")
        lines.append("")
    if h.score_explanations:
        lines.append("**Объяснение оценок:**")
        for text in h.score_explanations.values():
            lines.append(f"- {text}")
        lines.append("")
    if h.sources:
        lines.append("**Источники:**")
        for src in h.sources:
            lines.append(f"- {_format_source(src)}")
        lines.append("")
    if h.verification_steps:
        lines.append("**Верификация:**")
        for step in h.verification_steps:
            lines.append(f"- {step}")
        lines.append("")
    tech, econ = _format_risks(h)
    lines.extend([f"**Риски (тех.):** {tech}", f"**Риски (экон.):** {econ}", ""])
    return lines


def result_to_json(result: PipelineResult, constraints: str = "") -> str:
    payload = result.model_dump()
    payload["constraints"] = constraints
    payload["exported_at"] = datetime.now().isoformat()
    return json.dumps(payload, ensure_ascii=False, indent=2)


def result_to_csv(result: PipelineResult, constraints: str = "") -> str:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(
        [
            "rank",
            "title",
            "full_statement",
            "mechanism",
            "kpi_impact",
            "score_total",
            "novelty",
            "groundedness",
            "value",
            "risk",
            "sources",
            "verification_steps",
            "case_id",
            "kpi_goal",
            "constraints",
        ]
    )
    for i, h in enumerate(result.hypotheses, 1):
        s = h.scores
        sources = "; ".join(_format_source(src) for src in (h.sources or []))
        steps = "; ".join(h.verification_steps or [])
        writer.writerow(
            [
                i,
                h.title,
                h.full_statement,
                h.mechanism or "",
                h.kpi_impact or "",
                f"{s.total:.3f}" if s else "",
                f"{s.novelty:.3f}" if s else "",
                f"{s.groundedness:.3f}" if s else "",
                f"{s.value:.3f}" if s else "",
                f"{s.risk:.3f}" if s else "",
                sources,
                steps,
                result.case_id,
                result.kpi_goal,
                constraints,
            ]
        )
    return buf.getvalue()


def result_to_docx_bytes(result: PipelineResult, constraints: str = "") -> bytes:
    doc = Document()
    doc.add_heading(f"Фабрика гипотез — {result.case_name}", level=0)
    meta = doc.add_paragraph()
    meta.add_run(f"KPI: {result.kpi_goal}\n").bold = True
    meta.add_run(f"Ограничения: {constraints or '—'}\n")
    meta.add_run(f"Режим: {result.mode}\n")
    meta.add_run(f"Дата: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    for i, h in enumerate(result.hypotheses, 1):
        doc.add_heading(f"{i}. {h.title}", level=1)
        p = doc.add_paragraph()
        run = p.add_run(h.full_statement)
        run.italic = True
        if h.mechanism:
            doc.add_paragraph(f"Механизм: {h.mechanism}")
        if h.kpi_impact:
            doc.add_paragraph(f"Влияние на KPI: {h.kpi_impact}")
        if h.scores:
            s = h.scores
            doc.add_paragraph(
                f"Оценки: итого {s.total:.2f} | новизна {s.novelty:.2f} | "
                f"обоснованность {s.groundedness:.2f} | ценность {s.value:.2f} | риск {s.risk:.2f}"
            )
        if h.sources:
            doc.add_paragraph("Источники:", style="List Bullet")
            for src in h.sources:
                doc.add_paragraph(_format_source(src), style="List Bullet")
        if h.verification_steps:
            doc.add_paragraph("Верификация:", style="List Bullet")
            for step in h.verification_steps:
                doc.add_paragraph(step, style="List Bullet")
        tech, econ = _format_risks(h)
        doc.add_paragraph(f"Риски: техн. — {tech}; экон. — {econ}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def result_to_pdf_bytes(result: PipelineResult, constraints: str = "") -> bytes:
    font_name = _register_pdf_font()
    styles = _pdf_styles(font_name)
    story: list[Any] = [
        _pdf_paragraph(f"Фабрика гипотез — {result.case_name}", styles["title"]),
        _pdf_paragraph(f"KPI: {result.kpi_goal}", styles["meta"]),
        _pdf_paragraph(f"Ограничения: {constraints or '—'}", styles["meta"]),
        _pdf_paragraph(f"Режим: {result.mode}", styles["meta"]),
        _pdf_paragraph(f"Дата: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles["meta"]),
        Spacer(1, 4 * mm),
    ]

    for i, h in enumerate(result.hypotheses, 1):
        story.append(_pdf_paragraph(f"{i}. {h.title}", styles["heading"]))
        story.append(_pdf_paragraph(h.full_statement, styles["statement"]))
        if h.mechanism:
            story.append(_pdf_paragraph(f"Механизм: {h.mechanism}", styles["body"]))
        if h.kpi_impact:
            story.append(_pdf_paragraph(f"Влияние на KPI: {h.kpi_impact}", styles["body"]))
        if h.scores:
            s = h.scores
            story.append(
                _pdf_paragraph(
                    f"Оценки: итого {s.total:.2f} | новизна {s.novelty:.2f} | "
                    f"обоснованность {s.groundedness:.2f} | ценность {s.value:.2f} | риск {s.risk:.2f}",
                    styles["body"],
                )
            )
        if h.sources:
            story.append(_pdf_paragraph("Источники:", styles["body"]))
            for src in h.sources:
                story.append(_pdf_paragraph(f"• {_format_source(src)}", styles["body"]))
        if h.verification_steps:
            story.append(_pdf_paragraph("Верификация:", styles["body"]))
            for step in h.verification_steps:
                story.append(_pdf_paragraph(f"• {step}", styles["body"]))
        tech, econ = _format_risks(h)
        story.append(_pdf_paragraph(f"Риски: техн. — {tech}; экон. — {econ}", styles["body"]))
        story.append(Spacer(1, 3 * mm))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=15 * mm,
        rightMargin=15 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
        title=f"Hypotheses {result.case_id}",
        author="Hypothesis Factory",
    )
    doc.build(story)
    return buf.getvalue()
